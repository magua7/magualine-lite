from __future__ import annotations

import ast
import io
import json
import re
import secrets
import threading
import time
from datetime import datetime, timedelta, timezone
from ipaddress import ip_address, ip_network
from pathlib import Path

import uvicorn
from fastapi import FastAPI, HTTPException, Request
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.sessions import SessionMiddleware

from .agent_client import AgentCallError, call_agent
from .config import get_settings
from .ip_geo import lookup_ip_geo, should_cache_geo_result
from .storage import (
    add_auth_attempt,
    add_blocked_ip,
    bulk_update_log_status,
    cache_ip_geo,
    clear_recent_auth_failures,
    get_cached_ip_geo,
    get_agent_status_items,
    get_recent_auth_failure_state,
    list_cc_bans,
    get_ip_analysis_data,
    get_log_detail,
    get_overview,
    get_screen_detail_data,
    get_screen_data,
    get_screen_summary_data,
    init_db,
    list_blocked_ips,
    remove_cc_ban,
    list_logs,
    remove_blocked_ip,
    update_log_status,
)


settings = get_settings()
BASE_DIR = Path(__file__).resolve().parent
templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))

DISPOSITION_DISPLAY_LABELS = {
    "real_attack": "真实攻击行为",
    "customer_business": "客户业务行为",
    "pending_business": "待确认业务行为",
    "notified_event": "已通报事件告警",
    "whitelist_traffic": "白名单流量",
    "pending": "待确认业务行为",
    "resolved": "已通报事件告警",
    "resolved_event": "已通报事件告警",
    "not_applicable": "未分类流量",
}

RISK_LEVEL_DISPLAY_LABELS = {
    "critical": "严重",
    "high": "高危",
    "medium": "中危",
    "low": "低危",
}

ATTACK_TYPE_DISPLAY_LABELS = {
    "manual_block": "手动封禁",
    "sql_injection": "SQL 注入",
    "xss": "跨站脚本",
    "ssti": "模板注入",
    "ssrf": "服务端请求伪造",
    "xxe_injection": "XXE 实体注入",
    "xxe_dtd_subset_decl": "XXE DTD 子集声明",
    "xxe_entity_declaration": "XXE 实体声明",
    "xxe_external_entity_uri": "XXE 外部实体引用",
    "nosql_injection": "NoSQL 注入",
    "nosql_operator_payload": "NoSQL 操作符注入",
    "nosql_regex_payload": "NoSQL 正则注入",
    "nosql_where_javascript": "NoSQL 脚本注入",
    "ldap_injection": "LDAP 注入",
    "ldap_wildcard_auth_bypass": "LDAP 通配绕过",
    "ldap_boolean_injection": "LDAP 布尔注入",
    "ldap_objectclass_enumeration": "LDAP 枚举探测",
    "file_inclusion": "文件包含",
    "file_inclusion_stream_wrapper": "文件包含流包装器",
    "file_inclusion_remote_url_param": "远程文件包含",
    "file_inclusion_local_target": "本地文件泄露探测",
    "path_traversal": "目录穿越",
    "command_injection": "命令注入",
    "deserialization_probe": "反序列化探测",
    "scanner_probe": "扫描探测",
    "scanner_probe_extended": "扫描器指纹探测",
    "scanner_probe_recon_suite": "侦察工具指纹",
    "scanner_probe_oast_marker": "OAST 探测标记",
    "scanner_probe_fuzz_placeholder": "模糊测试占位符",
    "sensitive_probe": "敏感路径探测",
    "sensitive_probe_extended": "敏感路径探测",
    "sensitive_probe_repository_metadata": "仓库元数据探测",
    "sensitive_probe_admin_interfaces": "管理界面探测",
    "sensitive_probe_config_leak": "配置文件泄露探测",
    "sensitive_probe_debug_surface": "调试面探测",
    "java_ecosystem_probe": "Java 生态路径探测",
    "java_probe_actuator_surface": "Spring Actuator 探测",
    "java_probe_middleware_surface": "Java 中间件控制台探测",
    "java_probe_archive_descriptor": "Java 描述文件探测",
    "php_ecosystem_probe": "PHP 生态路径探测",
    "php_probe_debug_surface": "PHP 调试端点探测",
    "php_probe_dependency_surface": "PHP 依赖目录探测",
    "php_probe_storage_surface": "PHP 存储日志探测",
    "brute_force": "暴力破解",
    "webshell_upload": "WebShell 上传",
    "webshell_probe": "WebShell 探测",
    "cve_exploit_attempt": "CVE 漏洞利用",
    "cve_log4shell": "Log4Shell 利用",
    "cve_spring4shell": "Spring4Shell 利用",
    "cve_struts_ognl": "Struts OGNL 利用",
    "cve_confluence_ognl": "Confluence OGNL 利用",
    "cve_citrix_traversal": "Citrix 路径穿越",
    "cve_apache_traversal": "Apache 路径穿越",
    "cve_phpunit_eval_stdin": "PHPUnit 利用",
    "cve_thinkphp_rce": "ThinkPHP 远程执行",
    "cve_fastjson_auto_type": "Fastjson 利用",
    "cve_weblogic_console_traversal": "WebLogic 利用",
    "cve_jboss_invoker_deserialization": "JBoss 利用",
    "cve_spring_gateway_spel": "Spring Gateway 利用",
    "cve_laravel_ignition_rce": "Laravel Ignition 利用",
    "cve_php_cgi_arg_injection": "PHP CGI 利用",
    "cve_drupalgeddon2": "Drupal 利用",
    "cve_yii_debug_rce": "Yii Debug 利用",
}

app = FastAPI(docs_url=None, redoc_url=None, openapi_url=None, title="magualine-admin")
app.add_middleware(SessionMiddleware, secret_key=settings.secret_key, same_site="lax")
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

LOG_ANALYSIS_SECTION_KEYS = (
    "title",
    "overview_summary",
    "threat_level",
    "anomaly_ips",
    "key_findings",
    "attack_patterns",
    "risk_summary",
    "recommendations",
    "security_notices",
    "ban_policy_advice",
    "next_steps",
)

SCREEN_CACHE_TTL_SECONDS = 7
SCREEN_CACHE_HOURS = 24
ADMIN_LOGIN_PATH = "/api/login"
_screen_cache_lock = threading.Lock()
_screen_cache = {
    "hours": None,
    "generated_at": "",
    "expires_at": 0.0,
    "payload": None,
}


def _parse_ip_value(value: str | None) -> str | None:
    candidate = str(value or "").strip()
    if not candidate:
        return None
    try:
        return str(ip_address(candidate))
    except ValueError:
        return None


def _extract_forwarded_for_ip(value: str | None) -> str | None:
    if not value:
        return None
    for item in str(value).split(","):
        candidate = _parse_ip_value(item)
        if candidate:
            return candidate
    return None


def _is_trusted_proxy_ip(value: str | None) -> bool:
    parsed = _parse_ip_value(value)
    if not parsed:
        return False

    source_ip = ip_address(parsed)
    for proxy in settings.trusted_proxy_ips:
        trusted_proxy = str(proxy or "").strip()
        if not trusted_proxy:
            continue
        try:
            if "/" in trusted_proxy:
                if source_ip in ip_network(trusted_proxy, strict=False):
                    return True
            elif source_ip == ip_address(trusted_proxy):
                return True
        except ValueError:
            continue
    return False


def get_request_client_ip(request: Request) -> str:
    source_ip = request.client.host if request.client else "unknown"
    if not _is_trusted_proxy_ip(source_ip):
        return source_ip

    forwarded_for = _extract_forwarded_for_ip(request.headers.get("x-forwarded-for"))
    if forwarded_for:
        return forwarded_for

    real_ip = _parse_ip_value(request.headers.get("x-real-ip"))
    if real_ip:
        return real_ip

    return source_ip


def get_admin_login_lock_state(client_ip: str) -> dict[str, object]:
    failure_state = get_recent_auth_failure_state(
        client_ip,
        path=ADMIN_LOGIN_PATH,
        window_seconds=settings.admin_login_window_seconds,
    )
    failed_count = int(failure_state.get("count") or 0)
    locked_until = ""
    retry_after = 0

    if failed_count >= settings.admin_login_max_failures:
        try:
            last_failed_at = datetime.fromisoformat(str(failure_state.get("last_failed_at") or "")).astimezone(timezone.utc)
            expires_at = last_failed_at + timedelta(seconds=settings.admin_login_lock_seconds)
            retry_after = max(0, int((expires_at - datetime.now(timezone.utc)).total_seconds()))
            if retry_after > 0:
                locked_until = expires_at.isoformat()
        except Exception:
            retry_after = settings.admin_login_lock_seconds

    return {
        "failed_count": failed_count,
        "retry_after": retry_after,
        "locked": retry_after > 0,
        "locked_until": locked_until,
    }


def get_cached_screen_data(hours: int = SCREEN_CACHE_HOURS) -> dict:
    return get_screen_data(hours=hours)

LOG_ANALYSIS_SECTION_PATTERN = re.compile(
    r'(?<![\w/])["\']?(title|overview_summary|threat_level|anomaly_ips|key_findings|attack_patterns|risk_summary|recommendations|security_notices|ban_policy_advice|next_steps)["\']?\s*[:：]',
    re.IGNORECASE,
)

ANOMALY_IP_FIELD_PATTERN = re.compile(r'["\']?ip["\']?\s*[:：]\s*["\']?((?:\d{1,3}\.){3}\d{1,3})["\']?', re.IGNORECASE)


def _to_compact_json(data: object) -> str:
    return json.dumps(data, ensure_ascii=False, separators=(",", ":"))


def _truncate_text(value: object, limit: int = 600) -> str:
    text = str(value or "").strip()
    if len(text) <= limit:
        return text
    return f"{text[:limit]}...(已截断)"


def _prepare_overview_for_agent(overview: dict) -> dict:
    compact = dict(overview)
    compact["latest_high_risk_alerts"] = (overview.get("latest_high_risk_alerts") or [])[:5]
    compact["recent_alert_stream"] = (overview.get("recent_alert_stream") or [])[:8]
    compact["top_source_ips"] = (overview.get("top_source_ips") or [])[:8]
    compact["top_attack_types"] = (overview.get("top_attack_types") or [])[:8]
    compact["top_paths"] = (overview.get("top_paths") or [])[:8]
    compact["geo_buckets"] = (overview.get("geo_buckets") or [])[:8]
    compact["hourly_trend"] = (overview.get("hourly_trend") or [])[-12:]
    return compact


def _prepare_log_detail_for_agent(log_detail: dict) -> dict:
    headers = log_detail.get("request_headers")
    if isinstance(headers, dict):
        compact_headers = {}
        for key in (
            "host",
            "user-agent",
            "content-type",
            "referer",
            "origin",
            "x-forwarded-for",
            "authorization",
            "cookie",
        ):
            value = headers.get(key)
            if value is not None:
                compact_headers[key] = _truncate_text(value, 180)
        headers = compact_headers

    return {
        "id": log_detail.get("id"),
        "created_at": log_detail.get("created_at"),
        "client_ip": log_detail.get("client_ip"),
        "destination_host": log_detail.get("destination_host"),
        "destination_ip": log_detail.get("destination_ip"),
        "ip_geo": log_detail.get("ip_geo", {}),
        "method": log_detail.get("method"),
        "path": log_detail.get("path"),
        "query_string": _truncate_text(log_detail.get("query_string"), 300),
        "action": log_detail.get("action"),
        "attack_type": log_detail.get("attack_type"),
        "attack_detail": _truncate_text(log_detail.get("attack_detail"), 500),
        "cve_id": log_detail.get("cve_id"),
        "severity": log_detail.get("severity"),
        "alert_status": log_detail.get("alert_status"),
        "handled_status": log_detail.get("handled_status"),
        "status_code": log_detail.get("status_code"),
        "upstream_status": log_detail.get("upstream_status"),
        "duration_ms": log_detail.get("duration_ms"),
        "request_headers": headers or {},
        "body_preview": _truncate_text(log_detail.get("body_preview"), 1200),
    }


def _build_overview_prompt(overview: dict, alert_samples: list[dict], blocked_ips: list[dict]) -> str:
    return (
        "任务模式：overview_24h\n"
        "请基于输入的数据完成过去 24 小时安全态势分析，输出结构化 JSON 格式的分析报告。\n\n"
        "字段值必须直接面向终端用户，使用简洁清晰的中文自然语言；"
        "不要输出 Markdown 代码块、不要输出 JSON 解释、不要在文本里夹带字段名或程序片段。\n\n"
        "必须包含以下字段（JSON 格式）：\n"
        "{\n"
        '  "title": "24小时安全态势分析",\n'
        '  "summary": "整体态势摘要",\n'
        '  "key_findings": ["关键发现1", "关键发现2", ...],\n'
        '  "actions_now": ["立即采取的行动1", "立即采取的行动2", ...],\n'
        '  "actions_today": ["今日需要完成的行动1", "今日需要完成的行动2", ...],\n'
        '  "watch_list": ["需要持续关注的项目1", "需要持续关注的项目2", ...],\n'
        '  "rule_improvement_directions": ["规则改进方向1", "规则改进方向2", ...],\n'
        '  "false_positive_risks": ["误报风险1", "误报风险2", ...],\n'
        '  "confidence": "0.0-1.0 之间的置信度数值"\n'
        "}\n\n"
        f"过去24小时概览数据：\n{_to_compact_json(_prepare_overview_for_agent(overview))}\n\n"
        f"高危告警样本（最多5条）：\n{_to_compact_json(alert_samples)}\n\n"
        f"当前封禁IP列表（最多12条）：\n{_to_compact_json(blocked_ips)}\n"
    )


def _build_single_log_prompt(log_detail: dict) -> str:
    return (
        "任务模式：single_flow_triage\n"
        "请对单条流量进行研判，输出结构化 JSON 格式的分析报告。\n\n"
        "所有字符串字段都必须直接面向安全运营人员，用中文自然语言表达；"
        "不要输出 Markdown 代码块、不要输出程序语言、不要在字段值里重复 JSON 键名。\n\n"
        "必须包含以下字段（JSON 格式）：\n"
        "{\n"
        '  "title": "简短的研判结论标题",\n'
        '  "disposition": "real_attack|customer_business|pending_business|notified_event",\n'
        '  "risk_level": "critical|high|medium|low",\n'
        '  "confidence": "0.0-1.0 之间的置信度数值",\n'
        '  "attack_analysis": "详细的攻击分析说明",\n'
        '  "evidence": ["证据1", "证据2", ...],\n'
        '  "uncertainties": ["不确定点1", "不确定点2", ...],\n'
        '  "suggested_actions": ["建议动作1", "建议动作2", ...],\n'
        '  "rule_patch_suggestion": [{"type": "...", "target": "...", "proposal": "..."}, ...]\n'
        "}\n\n"
        "分类标准：\n"
        "- real_attack: 确认为真实攻击行为（漏洞利用、暴力破解、恶意扫描等）\n"
        "- customer_business: 命中规则但属于客户正常业务\n"
        "- pending_business: 无法直接确定真伪，需要继续观察\n"
        "- notified_event: 已完成研判和通报的事件\n\n"
        f"流量详情：\n{_to_compact_json(_prepare_log_detail_for_agent(log_detail))}\n"
    )


def _build_log_analysis_prompt(analysis_data: dict) -> str:
    summary = analysis_data.get("summary", {})
    return (
        "任务模式：log_analysis_24h\n"
        "请基于以下过去24小时的日志数据，对异常IP进行深度分析，输出结构化 JSON 格式的分析报告。\n\n"
        "所有字符串字段必须是面向用户的中文自然语言，不要在文本里混入 JSON 字段名、代码块、转义符或程序片段。\n\n"
        "必须包含以下字段（JSON 格式）：\n"
        "{\n"
        '  "title": "24小时日志异常IP分析报告",\n'
        '  "overview_summary": "整体流量和异常情况的简要描述",\n'
        '  "threat_level": "critical|high|medium|low",\n'
        '  "anomaly_ips": [\n'
        '    {\n'
        '      "ip": "IP地址",\n'
        '      "risk_level": "critical|high|medium|low",\n'
        '      "anomaly_type": "异常类型描述",\n'
        '      "evidence": "异常证据",\n'
        '      "ban_suggestion": {"duration_days": 封禁天数, "reason": "封禁原因"},\n'
        '      "priority": 处置优先级数字(1最高)\n'
        '    }\n'
        '  ],\n'
        '  "key_findings": ["关键发现1", "关键发现2", ...],\n'
        '  "attack_patterns": ["攻击模式分析1", "攻击模式分析2", ...],\n'
        '  "risk_summary": "整体风险评估说明",\n'
        '  "recommendations": ["安全建议1", "安全建议2", ...],\n'
        '  "security_notices": ["安全注意事项1", "安全注意事项2", ...],\n'
        '  "ban_policy_advice": "IP封禁策略总体建议",\n'
        '  "next_steps": ["下一步行动1", "下一步行动2", ...]\n'
        "}\n\n"
        "分析维度说明：\n"
        "1. 短时间内多次访问的IP：高频请求、疑似自动化攻击工具\n"
        "2. 访问量最多的IP：流量占比异常高、持续骚扰\n"
        "3. 攻击程度最严重的IP：命中高危规则、CVE利用、WebShell上传等\n"
        "4. 扫描探测类IP：访问大量不同路径、信息收集行为\n"
        "封禁时效建议维度：攻击频率、攻击类型严重程度、是否已知恶意IP段\n"
        "- 一般扫描探测：建议封禁 1-3 天\n"
        "- 暴力破解/SQL注入：建议封禁 7-14 天\n"
        "- WebShell上传/CVE利用/命令注入：建议封禁 30-90 天或永久\n\n"
        f"过去{analysis_data.get('window_hours', 24)}小时整体统计：\n"
        f"{_to_compact_json(summary)}\n\n"
        f"高频短时访问IP（访问次数>=10）：\n"
        f"{_to_compact_json((analysis_data.get('frequent_short_ips') or [])[:10])}\n\n"
        f"访问量最多IP（Top10）：\n"
        f"{_to_compact_json((analysis_data.get('top_access_ips') or [])[:10])}\n\n"
        f"攻击最严重IP（高危命中Top10）：\n"
        f"{_to_compact_json((analysis_data.get('most_dangerous_ips') or [])[:10])}\n\n"
        f"扫描探测IP（访问不同路径>=5）：\n"
        f"{_to_compact_json((analysis_data.get('scanner_ips') or [])[:10])}\n\n"
        f"当前已封禁IP列表：\n"
        f"{_to_compact_json((analysis_data.get('blocked_ips') or [])[:15])}\n"
    )


def _build_log_analysis_repair_prompt(raw_text: str) -> str:
    content = _truncate_text(raw_text, 6000)
    return (
        "任务模式：log_analysis_json_repair\n"
        "你会收到一段日志分析文本。请将其整理为严格 JSON 对象，并且只输出 JSON，不要输出其他说明。\n"
        "字段值必须是可直接展示给用户的中文自然语言，不要保留代码块、字段名解释或程序化格式。\n"
        "必须包含字段：title, overview_summary, threat_level, anomaly_ips, key_findings, attack_patterns, "
        "risk_summary, recommendations, security_notices, ban_policy_advice, next_steps。\n"
        "threat_level 只允许：critical|high|medium|low。\n"
        "anomaly_ips 必须是数组，元素包含：ip, risk_level, anomaly_type, evidence, ban_suggestion, priority。\n\n"
        f"待整理文本：\n{content}\n"
    )


def _build_log_analysis_display(parsed: dict, raw_text: str, analysis_data: dict) -> dict:
    text_hints = _build_log_analysis_text_hints(raw_text)

    if not parsed:
        return _ensure_log_analysis_display_sections(
            {
                "title": "AI 日志分析",
                "overview_summary": text_hints.get("overview_summary") or "模型未返回可直接展示的自然语言分析结果",
                "threat_level": text_hints.get("threat_level") or "medium",
                "anomaly_ips": text_hints.get("anomaly_ips") or [],
                "key_findings": text_hints.get("key_findings") or [],
                "attack_patterns": text_hints.get("attack_patterns") or [],
                "risk_summary": text_hints.get("risk_summary") or "",
                "recommendations": text_hints.get("recommendations") or [],
                "security_notices": text_hints.get("security_notices") or [],
                "ban_policy_advice": text_hints.get("ban_policy_advice") or "",
                "next_steps": text_hints.get("next_steps") or [],
                "raw_analysis_data": analysis_data,
                "raw_text": raw_text,
            }
        )

    display = {
        "title": _normalize_log_analysis_text_field(parsed.get("title"), "title", text_hints.get("title")) or "AI 日志分析",
        "overview_summary": _normalize_log_analysis_text_field(
            parsed.get("overview_summary") or parsed.get("summary"),
            "overview_summary",
            text_hints.get("overview_summary") or raw_text,
        ),
        "threat_level": _normalize_risk_level_key(parsed.get("threat_level") or text_hints.get("threat_level") or "medium"),
        "anomaly_ips": _normalize_anomaly_ips(parsed.get("anomaly_ips")) or text_hints.get("anomaly_ips") or [],
        "key_findings": _normalize_log_analysis_list_field(parsed.get("key_findings"), text_hints.get("key_findings")),
        "attack_patterns": _normalize_log_analysis_list_field(parsed.get("attack_patterns"), text_hints.get("attack_patterns")),
        "risk_summary": _normalize_log_analysis_text_field(
            parsed.get("risk_summary"),
            "risk_summary",
            text_hints.get("risk_summary"),
        ),
        "recommendations": _normalize_log_analysis_list_field(parsed.get("recommendations"), text_hints.get("recommendations")),
        "security_notices": _normalize_log_analysis_list_field(parsed.get("security_notices"), text_hints.get("security_notices")),
        "ban_policy_advice": _normalize_log_analysis_text_field(
            parsed.get("ban_policy_advice"),
            "ban_policy_advice",
            text_hints.get("ban_policy_advice"),
        ),
        "next_steps": _normalize_log_analysis_list_field(parsed.get("next_steps"), text_hints.get("next_steps")),
        "raw_analysis_data": analysis_data,
        "raw_text": raw_text,
    }
    return _ensure_log_analysis_display_sections(display)


def _build_local_log_analysis_display(analysis_data: dict, reason: str = "") -> dict:
    summary = analysis_data.get("summary", {}) if isinstance(analysis_data, dict) else {}
    anomaly_ips = _derive_local_anomaly_ips(analysis_data)

    high_count = int(summary.get("high_count") or 0)
    blocked_count = int(summary.get("blocked_count") or 0)
    threat_level = "medium"
    if high_count >= 20:
        threat_level = "critical"
    elif high_count >= 8 or blocked_count >= 100:
        threat_level = "high"
    elif high_count >= 3 or blocked_count >= 30:
        threat_level = "medium"
    else:
        threat_level = "low"

    key_findings = [
        f"24h 共接收 {int(summary.get('total_requests') or 0)} 次请求，涉及 {int(summary.get('unique_ips') or 0)} 个源 IP。",
        f"其中拦截 {blocked_count} 次，高危事件 {high_count} 次。",
    ]
    if anomaly_ips:
        key_findings.append(f"共识别 {len(anomaly_ips)} 个异常 IP，已按风险等级生成处置优先级。")

    notes = []
    if reason:
        notes.append(reason)

    return {
        "title": "24小时日志异常IP分析报告（降级模式）",
        "overview_summary": "AI 服务暂不可用，已基于本地日志规则引擎生成分析结果，可先用于应急研判与处置。",
        "threat_level": threat_level,
        "anomaly_ips": anomaly_ips,
        "key_findings": key_findings,
        "attack_patterns": [
            "短时高频访问与路径扫描并存，疑似自动化探测流量。",
            "部分源 IP 命中高危规则，需要优先封禁并复核目标接口。",
        ],
        "risk_summary": "当前结果来自本地规则聚合，不包含大模型语义研判，请在 AI 服务恢复后复跑确认。",
        "recommendations": [
            "优先处置 P1/P2 异常 IP，按建议时长执行封禁。",
            "针对高频命中路径补充限速、验证码与 WAF 精细规则。",
        ],
        "security_notices": notes,
        "ban_policy_advice": "扫描探测建议 1-3 天；暴力破解/注入建议 7-14 天；高危漏洞利用建议 30 天以上。",
        "next_steps": [
            "恢复 AI 配置后重新执行“立即分析”以获得语义化结论。",
            "导出当前报告并同步运维值班人员进行封禁落地。",
        ],
        "raw_analysis_data": analysis_data,
    }


def _strip_markdown_fence(text: str) -> str:
    fenced = re.match(r"^\s*```(?:json|markdown|md|text)?\s*(.*?)\s*```\s*$", text, re.DOTALL | re.IGNORECASE)
    if fenced:
        return fenced.group(1).strip()
    return text.strip()


def _coerce_json_like(value: object) -> object:
    if isinstance(value, (dict, list)):
        return value

    text = str(value or "").strip()
    if not text:
        return value

    cleaned = _strip_markdown_fence(text)
    cleaned = cleaned.replace("“", '"').replace("”", '"').replace("‘", "'").replace("’", "'").strip()
    cleaned = re.sub(r"^json\s*[:：]?\s*", "", cleaned, count=1, flags=re.IGNORECASE)
    candidates = [cleaned]

    for opener, closer in (("{", "}"), ("[", "]")):
        start_idx = cleaned.find(opener)
        end_idx = cleaned.rfind(closer)
        if start_idx != -1 and end_idx != -1 and start_idx < end_idx:
            candidates.append(cleaned[start_idx : end_idx + 1])

    for candidate in candidates:
        try:
            return json.loads(candidate)
        except Exception:
            pass
        try:
            return ast.literal_eval(candidate)
        except Exception:
            pass

    return value


def _extract_log_analysis_sections(raw_text: object) -> dict[str, str]:
    text = _strip_markdown_fence(str(raw_text or "").strip())
    if not text:
        return {}

    normalized = text.replace("\\n", "\n").replace("\r", "\n")
    normalized = re.sub(r"^json\s*[:：]?\s*", "", normalized, count=1, flags=re.IGNORECASE).strip()
    matches = list(LOG_ANALYSIS_SECTION_PATTERN.finditer(normalized))
    if not matches:
        return {}

    sections: dict[str, str] = {}
    for index, match in enumerate(matches):
        key = match.group(1).lower()
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(normalized)
        value = normalized[start:end].strip().strip(",").strip()
        sections[key] = value
    return sections


def _looks_like_log_analysis_field_dump(value: object, current_field: str | None = None) -> bool:
    text = _strip_markdown_fence(str(value or "").strip())
    if not text:
        return False

    current_key = (current_field or "").lower().strip()
    for match in LOG_ANALYSIS_SECTION_PATTERN.finditer(text):
        field = match.group(1).lower()
        if not current_key:
            return True
        if field != current_key:
            return True
        if match.start() == 0:
            return True
    return False


def _extract_named_text(block: str, field_name: str, next_fields: tuple[str, ...]) -> str:
    if not block:
        return ""
    lookahead = "|".join(re.escape(name) for name in next_fields)
    pattern = re.compile(
        rf'["\']?{re.escape(field_name)}["\']?\s*[:：]\s*(.+?)(?=(?:["\']?(?:{lookahead})["\']?\s*[:：])|$)',
        re.IGNORECASE | re.DOTALL,
    )
    match = pattern.search(block)
    if not match:
        return ""
    return match.group(1).strip().strip(",").strip()


def _extract_anomaly_ips_from_text(value: object) -> list[dict]:
    text = _strip_markdown_fence(str(value or "").strip())
    if not text:
        return []

    matches = list(ANOMALY_IP_FIELD_PATTERN.finditer(text))
    if not matches:
        return []

    fields = (
        "ip",
        "risk_level",
        "severity",
        "anomaly_type",
        "attack_type",
        "type",
        "evidence",
        "reason",
        "summary",
        "ban_suggestion",
        "duration_days",
        "days",
        "priority",
    )
    items: list[dict] = []
    for index, match in enumerate(matches):
        start = match.start()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        block = text[start:end]
        ip = match.group(1)
        risk_level = _normalize_risk_level_key(
            _extract_named_text(block, "risk_level", fields) or _extract_named_text(block, "severity", fields) or "medium"
        )
        anomaly_type = _humanize_attack_type(
            _extract_named_text(block, "anomaly_type", fields)
            or _extract_named_text(block, "attack_type", fields)
            or _extract_named_text(block, "type", fields)
        ) or "异常行为"
        evidence = _normalize_paragraph_text(
            _extract_named_text(block, "evidence", fields)
            or _extract_named_text(block, "reason", fields)
            or _extract_named_text(block, "summary", fields)
        )
        duration_text = _extract_named_text(block, "duration_days", fields) or _extract_named_text(block, "days", fields)
        try:
            duration_days = int(float(str(duration_text or "").strip()))
        except Exception:
            duration_days = _default_ban_suggestion(risk_level, anomaly_type).get("duration_days", 3)
        reason = _normalize_paragraph_text(_extract_named_text(block, "reason", fields)) or anomaly_type
        priority = _normalize_paragraph_text(_extract_named_text(block, "priority", fields))
        items.append(
            {
                "ip": ip,
                "risk_level": risk_level,
                "anomaly_type": anomaly_type,
                "evidence": evidence or f"检测到 {anomaly_type}",
                "ban_suggestion": {
                    "duration_days": duration_days,
                    "reason": reason,
                },
                "priority": priority,
            }
        )

    return _normalize_anomaly_ips(items)


def _normalize_log_analysis_text_field(primary: object, field_name: str, fallback: object = None) -> str:
    for candidate in (primary, fallback):
        text = _normalize_paragraph_text(candidate)
        if text and not _looks_like_log_analysis_field_dump(text, field_name):
            return text
    return ""


def _normalize_log_analysis_list_field(primary: object, fallback: object = None) -> list[str]:
    primary_items = [item for item in _normalize_list(primary) if not _looks_like_log_analysis_field_dump(item)]
    if primary_items:
        return primary_items
    return [item for item in _normalize_list(fallback) if not _looks_like_log_analysis_field_dump(item)]


def _build_log_analysis_text_hints(raw_text: str) -> dict:
    sections = _extract_log_analysis_sections(raw_text)
    anomaly_raw = sections.get("anomaly_ips") or raw_text
    anomaly_ips = _normalize_anomaly_ips(sections.get("anomaly_ips"))
    if not anomaly_ips:
        anomaly_ips = _extract_anomaly_ips_from_text(anomaly_raw)

    return {
        "title": _normalize_log_analysis_text_field(sections.get("title"), "title"),
        "overview_summary": _normalize_log_analysis_text_field(sections.get("overview_summary"), "overview_summary"),
        "threat_level": _normalize_risk_level_key(sections.get("threat_level") or "medium"),
        "anomaly_ips": anomaly_ips,
        "key_findings": _normalize_log_analysis_list_field(sections.get("key_findings")),
        "attack_patterns": _normalize_log_analysis_list_field(sections.get("attack_patterns")),
        "risk_summary": _normalize_log_analysis_text_field(sections.get("risk_summary"), "risk_summary"),
        "recommendations": _normalize_log_analysis_list_field(sections.get("recommendations")),
        "security_notices": _normalize_log_analysis_list_field(sections.get("security_notices")),
        "ban_policy_advice": _normalize_log_analysis_text_field(sections.get("ban_policy_advice"), "ban_policy_advice"),
        "next_steps": _normalize_log_analysis_list_field(sections.get("next_steps")),
    }


def _humanize_attack_type(value: object) -> str:
    key = str(value or "").strip()
    return ATTACK_TYPE_DISPLAY_LABELS.get(key, key)


def _normalize_disposition_key(value: object) -> str:
    key = str(value or "").strip()
    reverse_map = {
        "真实攻击行为": "real_attack",
        "客户业务行为": "customer_business",
        "待确认业务行为": "pending_business",
        "已通报事件告警": "notified_event",
        "白名单流量": "whitelist_traffic",
        "未分类流量": "not_applicable",
    }
    return reverse_map.get(key, key or "pending_business")


def _humanize_disposition(value: object) -> str:
    key = _normalize_disposition_key(value)
    return DISPOSITION_DISPLAY_LABELS.get(key, key or "未分类流量")


def _normalize_risk_level_key(value: object) -> str:
    key = str(value or "").strip().lower()
    reverse_map = {
        "严重": "critical",
        "高危": "high",
        "中危": "medium",
        "低危": "low",
    }
    return reverse_map.get(key, key or "medium")


def _humanize_risk_level(value: object) -> str:
    key = _normalize_risk_level_key(value)
    return RISK_LEVEL_DISPLAY_LABELS.get(key, key or "中危")


def _clean_natural_line(line: str) -> str:
    text = str(line or "").strip().strip(",").strip()
    text = text.strip('"').strip("'")
    text = re.sub(r"^(?:title|summary|attack_analysis|overview_summary|risk_summary|analysis|conclusion|content|text|message|reason|proposal|suggestion|evidence)\s*[:：]\s*", "", text, flags=re.IGNORECASE)
    text = re.sub(r"^[\-\*\u2022]\s*", "", text)
    if not re.match(r"^(?:\d{1,3}\.){3}\d{1,3}(?:\b|$)", text):
        text = re.sub(r"^\d+(?:[\.\)、\-])\s+", "", text)
    return text.strip()


def _normalize_paragraph_text(value: object) -> str:
    coerced = _coerce_json_like(value)

    if isinstance(coerced, list):
        return "\n".join(_normalize_list(coerced))

    if isinstance(coerced, dict):
        preferred_texts: list[str] = []
        for key in (
            "summary",
            "attack_analysis",
            "overview_summary",
            "risk_summary",
            "analysis",
            "description",
            "reason",
            "proposal",
            "suggestion",
            "content",
            "text",
            "message",
        ):
            candidate = _normalize_paragraph_text(coerced.get(key))
            if candidate:
                preferred_texts.append(candidate)
        if preferred_texts:
            return "\n".join(dict.fromkeys(preferred_texts))

        fallback_texts: list[str] = []
        for item in coerced.values():
            candidate = _normalize_paragraph_text(item)
            if candidate:
                fallback_texts.append(candidate)
            if len(fallback_texts) >= 3:
                break
        return "\n".join(dict.fromkeys(fallback_texts))

    text = str(coerced or "").strip()
    if not text:
        return ""

    text = _strip_markdown_fence(text)
    text = text.replace("\\n", "\n").replace("\r", "\n")
    text = re.sub(r"^json\s*[:：]?\s*", "", text, count=1, flags=re.IGNORECASE)

    cleaned_lines: list[str] = []
    for raw_line in text.split("\n"):
        line = _clean_natural_line(raw_line)
        if not line:
            continue
        if re.fullmatch(r"[\{\}\[\],]+", line):
            continue
        if re.fullmatch(r'"?[\w\u4e00-\u9fff\-\s]+"?\s*:\s*', line):
            continue
        if line.lower() in {"json", "output", "display"}:
            continue
        cleaned_lines.append(line)

    result_lines: list[str] = []
    for line in cleaned_lines:
        if line not in result_lines:
            result_lines.append(line)
    return "\n".join(result_lines).strip()


def _normalize_structured_note(item: dict) -> str:
    if not isinstance(item, dict):
        return _normalize_paragraph_text(item)

    for key in (
        "summary",
        "analysis",
        "description",
        "message",
        "reason",
        "proposal",
        "suggestion",
        "content",
        "text",
        "evidence",
        "attack_analysis",
    ):
        text = _normalize_paragraph_text(item.get(key))
        if text:
            return text

    parts: list[str] = []
    for key, value in item.items():
        text = _normalize_paragraph_text(value)
        if not text:
            continue
        if key in {"ip", "client_ip", "source_ip"}:
            parts.append(f"IP {text}")
        elif key in {"attack_type", "type", "category"}:
            parts.append(_humanize_attack_type(text))
        else:
            parts.append(text)
        if len(parts) >= 4:
            break
    return "；".join(parts)


def _normalize_list(value: object) -> list[str]:
    coerced = _coerce_json_like(value)

    if isinstance(coerced, list):
        raw_items = coerced
    elif isinstance(coerced, dict):
        raw_items = list(coerced.values())
    elif isinstance(coerced, str) and coerced.strip():
        if "\n" in coerced or "；" in coerced or ";" in coerced:
            raw_items = re.split(r"(?:\r?\n|；|;)+", coerced)
        else:
            raw_items = [coerced]
    else:
        raw_items = []

    result: list[str] = []
    for item in raw_items:
        parsed_item = _coerce_json_like(item)
        if isinstance(parsed_item, list):
            for nested in _normalize_list(parsed_item):
                if nested and nested not in result:
                    result.append(nested)
            continue

        if isinstance(parsed_item, dict):
            text = _normalize_structured_note(parsed_item)
        else:
            text = _normalize_paragraph_text(parsed_item)

        if not text:
            continue

        if "\n" in text:
            for sub in [segment.strip() for segment in text.splitlines() if segment.strip()]:
                if sub and sub not in result:
                    result.append(sub)
            continue

        if text not in result:
            result.append(text)
    return result


def _normalize_confidence(value: object) -> str:
    text = _normalize_paragraph_text(value)
    if not text:
        return ""
    normalized = text.replace("%", "").strip()
    try:
        number = float(normalized)
    except Exception:
        return text

    if number <= 1:
        number *= 100
    return f"{round(number)}%"


def _format_rule_patch_suggestion(item: object) -> str:
    parsed = _coerce_json_like(item)
    if isinstance(parsed, dict):
        target = _normalize_paragraph_text(parsed.get("target") or parsed.get("scope") or parsed.get("path") or parsed.get("location"))
        rule_type = _normalize_paragraph_text(parsed.get("type") or parsed.get("rule_type") or parsed.get("category"))
        proposal = _normalize_paragraph_text(parsed.get("proposal") or parsed.get("suggestion") or parsed.get("action") or parsed.get("content") or parsed.get("description"))

        if target and rule_type and proposal:
            return f"建议在 {target} 补充 {rule_type} 防护：{proposal}"
        if target and proposal:
            return f"建议在 {target} 补充防护：{proposal}"
        if rule_type and proposal:
            return f"建议补充 {rule_type} 防护：{proposal}"
        if proposal:
            return proposal
        return _normalize_structured_note(parsed)

    return _normalize_paragraph_text(parsed)


def _normalize_rule_patch_suggestions(value: object) -> list[str]:
    parsed = _coerce_json_like(value)
    if isinstance(parsed, list):
        items = parsed
    elif isinstance(parsed, dict):
        items = [parsed]
    elif isinstance(parsed, str) and parsed.strip():
        items = [parsed]
    else:
        items = []

    result: list[str] = []
    for item in items:
        text = _format_rule_patch_suggestion(item)
        if text and text not in result:
            result.append(text)
    return result


def _default_ban_suggestion(risk_level: str, anomaly_type: str) -> dict:
    if risk_level == "critical":
        days = 30
    elif risk_level == "high":
        days = 14
    elif risk_level == "medium":
        days = 7
    else:
        days = 3
    return {"duration_days": days, "reason": anomaly_type or "异常行为"}


def _extract_ip_from_text(text: str) -> str:
    match = re.search(r"\b(?:\d{1,3}\.){3}\d{1,3}\b", text or "")
    return match.group(0) if match else ""


def _normalize_ip_value(value: object) -> str:
    raw_text = _strip_markdown_fence(str(value or "").strip())
    raw_ip = _extract_ip_from_text(raw_text)
    if raw_ip:
        return raw_ip

    normalized = _normalize_paragraph_text(value)
    normalized_ip = _extract_ip_from_text(normalized)
    if normalized_ip:
        return normalized_ip

    return normalized if re.fullmatch(r"(?:\d{1,3}\.){3}\d{1,3}", normalized or "") else ""


def _normalize_anomaly_ip_item(item: object) -> dict | None:
    parsed = _coerce_json_like(item)

    if isinstance(parsed, str):
        raw_text = _strip_markdown_fence(str(parsed or "").strip())
        text = _normalize_paragraph_text(parsed)
        ip = _extract_ip_from_text(raw_text) or _extract_ip_from_text(text)
        if not ip:
            return None
        return {
            "ip": ip,
            "risk_level": "medium",
            "anomaly_type": "异常行为",
            "evidence": text,
            "ban_suggestion": _default_ban_suggestion("medium", "异常行为"),
            "priority": "",
        }

    if not isinstance(parsed, dict):
        return None

    ip = _normalize_ip_value(parsed.get("ip") or parsed.get("client_ip") or parsed.get("source_ip"))
    ip = ip or _extract_ip_from_text(json.dumps(parsed, ensure_ascii=False))
    ip = ip or _extract_ip_from_text(_normalize_paragraph_text(parsed))
    if not ip:
        return None

    risk_level = _normalize_risk_level_key(parsed.get("risk_level") or parsed.get("severity") or "medium")
    if risk_level not in {"critical", "high", "medium", "low"}:
        risk_level = "medium"

    anomaly_type = _normalize_paragraph_text(
        parsed.get("anomaly_type") or parsed.get("type") or parsed.get("category") or parsed.get("attack_type")
    )
    anomaly_type = _humanize_attack_type(anomaly_type) or "异常行为"

    evidence = _normalize_paragraph_text(
        parsed.get("evidence") or parsed.get("reason") or parsed.get("summary") or parsed.get("description")
    )
    if not evidence:
        snippets: list[str] = []
        total_count = parsed.get("total_count")
        high_count = parsed.get("high_count")
        unique_paths = parsed.get("unique_paths")
        attack_types = _normalize_paragraph_text(parsed.get("attack_types"))
        if total_count:
            snippets.append(f"总请求 {total_count} 次")
        if high_count:
            snippets.append(f"高危命中 {high_count} 次")
        if unique_paths:
            snippets.append(f"涉及 {unique_paths} 个路径")
        if attack_types:
            snippets.append(f"攻击类型：{_humanize_attack_type(attack_types)}")
        evidence = "；".join(snippets)

    ban_raw = _coerce_json_like(parsed.get("ban_suggestion") or parsed.get("ban"))
    if isinstance(ban_raw, dict):
        ban_suggestion = {
            "duration_days": ban_raw.get("duration_days") or ban_raw.get("days") or "",
            "reason": _normalize_paragraph_text(ban_raw.get("reason") or ban_raw.get("proposal") or anomaly_type),
        }
    else:
        ban_text = _normalize_paragraph_text(ban_raw)
        ban_suggestion = {
            "duration_days": "",
            "reason": ban_text or anomaly_type,
        }

    default_ban = _default_ban_suggestion(risk_level, anomaly_type)
    if not ban_suggestion.get("duration_days"):
        ban_suggestion["duration_days"] = default_ban["duration_days"]
    if not ban_suggestion.get("reason"):
        ban_suggestion["reason"] = default_ban["reason"]

    return {
        "ip": ip,
        "risk_level": risk_level,
        "anomaly_type": anomaly_type,
        "evidence": evidence or "待补充异常证据",
        "ban_suggestion": ban_suggestion,
        "priority": parsed.get("priority") or "",
    }


def _normalize_anomaly_ips(value: object) -> list[dict]:
    parsed = _coerce_json_like(value)
    if not isinstance(parsed, list):
        return []

    result: list[dict] = []
    seen_ips: set[str] = set()
    for item in parsed:
        normalized = _normalize_anomaly_ip_item(item)
        if not normalized:
            continue
        ip = normalized.get("ip", "")
        if ip in seen_ips:
            continue
        seen_ips.add(ip)
        result.append(normalized)
    return result


def _summarize_ip_item(ip_item: dict) -> str:
    ip = str(ip_item.get("ip") or "未知IP")
    anomaly_type = str(ip_item.get("anomaly_type") or "异常行为")
    evidence = str(ip_item.get("evidence") or "")
    ban = ip_item.get("ban_suggestion") or {}
    days = ban.get("duration_days") if isinstance(ban, dict) else ""
    reason = ban.get("reason") if isinstance(ban, dict) else ""
    parts = [f"IP {ip} 存在{anomaly_type}"]
    if evidence:
        parts.append(f"证据：{evidence}")
    if days or reason:
        suggestion = f"建议封禁 {days} 天" if days else "建议封禁"
        if reason:
            suggestion += f"，原因：{reason}"
        parts.append(suggestion)
    return "；".join(parts) + "。"


def _derive_local_anomaly_ips(analysis_data: dict) -> list[dict]:
    frequent_short_ips = analysis_data.get("frequent_short_ips") or []
    most_dangerous_ips = analysis_data.get("most_dangerous_ips") or []
    scanner_ips = analysis_data.get("scanner_ips") or []

    anomaly_map: dict[str, dict] = {}

    def push_ip(ip: str, risk_level: str, anomaly_type: str, evidence: str, duration_days: int, priority: int) -> None:
        if not ip:
            return
        exists = anomaly_map.get(ip)
        candidate = {
            "ip": ip,
            "risk_level": risk_level,
            "anomaly_type": anomaly_type,
            "evidence": evidence,
            "ban_suggestion": {"duration_days": duration_days, "reason": anomaly_type},
            "priority": priority,
        }
        if not exists:
            anomaly_map[ip] = candidate
            return
        rank = {"critical": 4, "high": 3, "medium": 2, "low": 1}
        if rank.get(risk_level, 0) > rank.get(exists.get("risk_level", "low"), 0):
            anomaly_map[ip] = candidate

    for item in most_dangerous_ips[:5]:
        if not isinstance(item, dict):
            continue
        ip = str(item.get("client_ip") or "").strip()
        high_count = int(item.get("high_count") or 0)
        attack_types = _humanize_attack_type(item.get("attack_types") or "未知攻击")
        push_ip(
            ip=ip,
            risk_level="critical" if high_count >= 5 else "high",
            anomaly_type="高危攻击行为",
            evidence=f"高危命中 {high_count} 次，攻击类型：{attack_types}",
            duration_days=30 if high_count >= 5 else 14,
            priority=1 if high_count >= 5 else 2,
        )

    for item in frequent_short_ips[:5]:
        if not isinstance(item, dict):
            continue
        ip = str(item.get("client_ip") or "").strip()
        total_count = int(item.get("total_count") or 0)
        blocked_count = int(item.get("blocked_count") or 0)
        push_ip(
            ip=ip,
            risk_level="high" if blocked_count >= 3 else "medium",
            anomaly_type="短时高频访问",
            evidence=f"24 小时请求 {total_count} 次，拦截 {blocked_count} 次",
            duration_days=7 if blocked_count >= 3 else 3,
            priority=2 if blocked_count >= 3 else 3,
        )

    for item in scanner_ips[:5]:
        if not isinstance(item, dict):
            continue
        ip = str(item.get("client_ip") or "").strip()
        unique_paths = int(item.get("unique_paths") or 0)
        total_count = int(item.get("total_count") or 0)
        push_ip(
            ip=ip,
            risk_level="medium",
            anomaly_type="扫描探测行为",
            evidence=f"访问 {unique_paths} 个不同路径，共 {total_count} 次请求",
            duration_days=3,
            priority=3,
        )

    return sorted(
        anomaly_map.values(),
        key=lambda x: (x.get("priority", 99), -int(x.get("ban_suggestion", {}).get("duration_days") or 0)),
    )[:10]


def _merge_anomaly_ips(primary: list[dict], fallback: list[dict]) -> list[dict]:
    if not fallback:
        return primary

    fallback_map = {str(item.get("ip") or ""): item for item in fallback if str(item.get("ip") or "").strip()}
    merged: list[dict] = []
    seen_ips: set[str] = set()

    for item in primary:
        ip = str(item.get("ip") or "").strip()
        if not ip:
            continue
        fallback_item = fallback_map.get(ip, {})
        ban = item.get("ban_suggestion") if isinstance(item.get("ban_suggestion"), dict) else {}
        fallback_ban = fallback_item.get("ban_suggestion") if isinstance(fallback_item.get("ban_suggestion"), dict) else {}
        merged_item = {
            "ip": ip,
            "risk_level": item.get("risk_level") or fallback_item.get("risk_level") or "medium",
            "anomaly_type": item.get("anomaly_type") or fallback_item.get("anomaly_type") or "异常行为",
            "evidence": item.get("evidence") or fallback_item.get("evidence") or "待补充异常证据",
            "ban_suggestion": {
                "duration_days": ban.get("duration_days") or fallback_ban.get("duration_days") or _default_ban_suggestion(item.get("risk_level") or fallback_item.get("risk_level") or "medium", item.get("anomaly_type") or fallback_item.get("anomaly_type") or "异常行为")["duration_days"],
                "reason": ban.get("reason") or fallback_ban.get("reason") or item.get("anomaly_type") or fallback_item.get("anomaly_type") or "异常行为",
            },
            "priority": item.get("priority") or fallback_item.get("priority") or "",
        }
        merged.append(merged_item)
        seen_ips.add(ip)

    for item in fallback:
        ip = str(item.get("ip") or "").strip()
        if not ip or ip in seen_ips:
            continue
        merged.append(item)
    return merged


def _ensure_log_analysis_display_sections(display: dict) -> dict:
    raw_analysis_data = display.get("raw_analysis_data") if isinstance(display.get("raw_analysis_data"), dict) else {}
    raw_text = str(display.get("raw_text") or "")
    text_hints = _build_log_analysis_text_hints(raw_text) if raw_text else {}
    fallback_display = _build_local_log_analysis_display(raw_analysis_data) if raw_analysis_data else {}
    fallback_anomaly_ips = _normalize_anomaly_ips(fallback_display.get("anomaly_ips", [])) if fallback_display else []

    anomaly_ips = _normalize_anomaly_ips(display.get("anomaly_ips"))
    if not anomaly_ips:
        anomaly_ips = text_hints.get("anomaly_ips") or []
    if not anomaly_ips:
        anomaly_ips = _extract_anomaly_ips_from_text(display.get("anomaly_ips"))
    anomaly_ips = _merge_anomaly_ips(anomaly_ips, fallback_anomaly_ips) if fallback_anomaly_ips else anomaly_ips
    anomaly_ips = _normalize_anomaly_ips(anomaly_ips)

    title = _normalize_log_analysis_text_field(display.get("title"), "title", text_hints.get("title") or fallback_display.get("title"))
    overview_summary = _normalize_log_analysis_text_field(
        display.get("overview_summary"),
        "overview_summary",
        text_hints.get("overview_summary") or fallback_display.get("overview_summary"),
    )
    risk_summary = _normalize_log_analysis_text_field(
        display.get("risk_summary"),
        "risk_summary",
        text_hints.get("risk_summary") or fallback_display.get("risk_summary"),
    )
    ban_policy_advice = _normalize_log_analysis_text_field(
        display.get("ban_policy_advice"),
        "ban_policy_advice",
        text_hints.get("ban_policy_advice") or fallback_display.get("ban_policy_advice"),
    )

    key_findings = _normalize_log_analysis_list_field(display.get("key_findings"), text_hints.get("key_findings") or fallback_display.get("key_findings"))
    if not key_findings and anomaly_ips:
        key_findings = [_summarize_ip_item(item) for item in anomaly_ips[:3]]

    attack_patterns = _normalize_log_analysis_list_field(display.get("attack_patterns"), text_hints.get("attack_patterns") or fallback_display.get("attack_patterns"))
    if not attack_patterns and anomaly_ips:
        focus_ips = ", ".join([item["ip"] for item in anomaly_ips[:3] if item.get("ip")])
        if focus_ips:
            attack_patterns = [f"异常流量主要集中在 {focus_ips}。"]

    recommendations = _normalize_log_analysis_list_field(display.get("recommendations"), text_hints.get("recommendations") or fallback_display.get("recommendations"))
    if not recommendations and anomaly_ips:
        recommendations = ["优先封禁高风险异常 IP，并对对应访问路径进行加固。"]

    security_notices = _normalize_log_analysis_list_field(display.get("security_notices"), text_hints.get("security_notices") or fallback_display.get("security_notices"))
    next_steps = _normalize_log_analysis_list_field(display.get("next_steps"), text_hints.get("next_steps") or fallback_display.get("next_steps"))
    if not next_steps and anomaly_ips:
        next_steps = ["复核异常 IP 对应日志明细，并将高风险源加入持续观察名单。"]

    display["title"] = title or "24 小时日志异常 IP 分析报告"
    display["overview_summary"] = overview_summary or "暂无可直接展示的 AI 总结，已回退到本地规则聚合结果。"
    display["risk_summary"] = risk_summary
    display["ban_policy_advice"] = ban_policy_advice
    display["threat_level"] = _normalize_risk_level_key(display.get("threat_level") or fallback_display.get("threat_level") or "medium")
    display["anomaly_ips"] = anomaly_ips
    display["key_findings"] = key_findings
    display["attack_patterns"] = attack_patterns
    display["recommendations"] = recommendations
    display["security_notices"] = security_notices
    display["next_steps"] = next_steps
    display["raw_analysis_data"] = raw_analysis_data or fallback_display.get("raw_analysis_data") or {}
    display["raw_text"] = raw_text
    return display


def _build_overview_display(parsed: dict, raw_text: str) -> dict:
    if not parsed:
        return {
            "title": "AI 态势分析",
            "summary": _normalize_paragraph_text(raw_text) or "模型未返回结构化结果",
            "key_findings": [],
            "actions_now": [],
            "actions_today": [],
            "watch_list": [],
        }

    return {
        "title": _normalize_paragraph_text(parsed.get("title")) or "AI 态势分析",
        "summary": _normalize_paragraph_text(parsed.get("summary") or parsed.get("overall_summary") or raw_text),
        "key_findings": _normalize_list(parsed.get("key_findings")),
        "actions_now": _normalize_list(parsed.get("actions_now")),
        "actions_today": _normalize_list(parsed.get("actions_today")),
        "watch_list": _normalize_list(parsed.get("watch_list")),
        "rule_improvement_directions": _normalize_list(parsed.get("rule_improvement_directions")),
        "false_positive_risks": _normalize_list(parsed.get("false_positive_risks")),
        "confidence": parsed.get("confidence", ""),
    }


def _build_log_display(parsed: dict, raw_text: str) -> dict:
    if not parsed:
        return {
            "title": "AI 单条流量研判",
            "summary": _normalize_paragraph_text(raw_text) or "模型未返回结构化结果",
            "disposition": "pending_business",
            "disposition_label": _humanize_disposition("pending_business"),
            "risk_level": "medium",
            "risk_level_label": _humanize_risk_level("medium"),
            "confidence": "",
            "evidence": [],
            "uncertainties": [],
            "suggested_actions": [],
            "rule_patch_suggestion": [],
        }

    disposition = _normalize_disposition_key(parsed.get("disposition") or "pending_business")
    risk_level = _normalize_risk_level_key(parsed.get("risk_level") or "medium")
    disposition_label = _humanize_disposition(disposition)
    risk_level_label = _humanize_risk_level(risk_level)

    evidence = _normalize_list(parsed.get("evidence"))
    uncertainties = _normalize_list(parsed.get("uncertainties"))
    suggested_actions = _normalize_list(parsed.get("suggested_actions"))
    rule_patch_suggestion = _normalize_rule_patch_suggestions(parsed.get("rule_patch_suggestion", []))

    title = _normalize_paragraph_text(parsed.get("title")) or f"AI 研判：{disposition_label}"
    summary = _normalize_paragraph_text(parsed.get("attack_analysis") or parsed.get("summary") or raw_text)
    if not summary:
        summary = f"该流量初步判断为{disposition_label}，风险等级为{risk_level_label}。"
        if evidence:
            summary += f" 主要依据：{evidence[0]}"
        if suggested_actions:
            summary += f" 建议优先执行：{suggested_actions[0]}"

    return {
        "title": title,
        "summary": summary,
        "disposition": disposition,
        "disposition_label": disposition_label,
        "risk_level": risk_level,
        "risk_level_label": risk_level_label,
        "confidence": _normalize_confidence(parsed.get("confidence", "")),
        "evidence": evidence,
        "uncertainties": uncertainties,
        "suggested_actions": suggested_actions,
        "rule_patch_suggestion": rule_patch_suggestion,
    }


def is_authenticated(request: Request) -> bool:
    return bool(request.session.get("authenticated"))


def require_api_auth(request: Request) -> None:
    if not is_authenticated(request):
        raise HTTPException(status_code=401, detail="请先登录")


@app.on_event("startup")
async def startup() -> None:
    init_db()


@app.get("/health")
async def health() -> dict[str, str]:
    return {"status": "ok"}


@app.get("/", response_class=HTMLResponse)
async def root(request: Request):
    if is_authenticated(request):
        return RedirectResponse(url="/dashboard", status_code=302)
    return RedirectResponse(url="/login", status_code=302)


@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    if is_authenticated(request):
        return RedirectResponse(url="/dashboard", status_code=302)
    return templates.TemplateResponse(
        request=request,
        name="login.html",
        context={"app_name": "magualine"},
    )


@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard_page(request: Request):
    if not is_authenticated(request):
        return RedirectResponse(url="/login", status_code=302)
    return templates.TemplateResponse(
        request=request,
        name="dashboard.html",
        context={"app_name": "magualine", "active_page": "dashboard"},
    )


@app.get("/screen", response_class=HTMLResponse)
async def screen_page(request: Request):
    if not is_authenticated(request):
        return RedirectResponse(url="/login", status_code=302)
    return templates.TemplateResponse(
        request=request,
        name="screen.html",
        context={"app_name": "magualine", "active_page": "screen"},
    )


@app.get("/logs", response_class=HTMLResponse)
async def logs_page(request: Request):
    if not is_authenticated(request):
        return RedirectResponse(url="/login", status_code=302)
    return templates.TemplateResponse(
        request=request,
        name="logs.html",
        context={"app_name": "magualine", "active_page": "logs"},
    )


@app.get("/log-analysis", response_class=HTMLResponse)
async def log_analysis_page(request: Request):
    if not is_authenticated(request):
        return RedirectResponse(url="/login", status_code=302)
    return templates.TemplateResponse(
        request=request,
        name="log_analysis.html",
        context={"app_name": "magualine", "active_page": "log-analysis"},
    )


@app.get("/block", response_class=HTMLResponse)
async def block_page(request: Request):
    if not is_authenticated(request):
        return RedirectResponse(url="/login", status_code=302)
    return templates.TemplateResponse(
        request=request,
        name="block.html",
        context={"app_name": "magualine", "active_page": "block"},
    )


@app.post("/api/login")
async def login(request: Request):
    client_ip = get_request_client_ip(request)
    lock_state = get_admin_login_lock_state(client_ip)
    if lock_state["locked"]:
        return JSONResponse(
            status_code=429,
            content={"message": "登录失败次数过多，请稍后再试"},
            headers={"Retry-After": str(lock_state["retry_after"])},
        )

    try:
        payload = await request.json()
    except Exception:
        payload = {}
    username = str(payload.get("username", "")).strip()
    password = str(payload.get("password", "")).strip()

    auth_ok = secrets.compare_digest(username, settings.admin_username) and secrets.compare_digest(
        password,
        settings.admin_password,
    )
    if not auth_ok:
        add_auth_attempt(client_ip, ADMIN_LOGIN_PATH, False, 401)
        next_state = get_admin_login_lock_state(client_ip)
        headers = {"Retry-After": str(next_state["retry_after"])} if next_state["locked"] else None
        status_code = 429 if next_state["locked"] else 401
        message = "登录失败次数过多，请稍后再试" if next_state["locked"] else "用户名或密码错误"
        return JSONResponse(status_code=status_code, content={"message": message}, headers=headers)

    add_auth_attempt(client_ip, ADMIN_LOGIN_PATH, True, 200)
    clear_recent_auth_failures(client_ip, path=ADMIN_LOGIN_PATH)
    request.session.clear()
    request.session["authenticated"] = True
    request.session["username"] = username
    return {"message": "ok"}


@app.post("/api/logout")
async def logout(request: Request):
    request.session.clear()
    return {"message": "ok"}


@app.get("/api/runtime")
async def runtime(request: Request):
    require_api_auth(request)
    return {"app_name": "magualine", "username": request.session.get("username", "admin")}


@app.get("/api/overview")
async def overview(request: Request):
    require_api_auth(request)
    return get_overview(hours=24)


@app.get("/api/screen")
async def screen_data(request: Request):
    require_api_auth(request)
    return get_cached_screen_data(hours=SCREEN_CACHE_HOURS)


@app.get("/api/screen/summary")
async def screen_summary_data(request: Request):
    require_api_auth(request)
    return get_screen_summary_data(hours=SCREEN_CACHE_HOURS)


@app.get("/api/screen/detail")
async def screen_detail_data(request: Request):
    require_api_auth(request)
    return get_screen_detail_data(hours=SCREEN_CACHE_HOURS)


@app.get("/api/agent/status")
async def agent_status(request: Request):
    require_api_auth(request)
    items = get_agent_status_items()
    latest_seen = ""
    for item in items:
        candidate = str(item.get("last_seen") or "").strip()
        if candidate and candidate > latest_seen:
            latest_seen = candidate
    overall_status = "online" if any(str(item.get("status") or "").strip() == "online" for item in items) else "offline"
    return {
        "status": overall_status,
        "items": items,
        "last_seen": latest_seen,
    }


@app.post("/api/agent/overview-24h")
async def agent_overview(request: Request):
    require_api_auth(request)
    payload = await request.json() if request.headers.get("content-type", "").startswith("application/json") else {}
    session_id = str(payload.get("session_id", "")).strip() or None

    overview_data = get_overview(hours=24)
    alert_samples = list_logs(page=1, page_size=12, alerts_only=True).get("items", [])
    blocked_ips_payload = list_blocked_ips(page=1, page_size=12)
    blocked_ips = blocked_ips_payload.get("items", [])
    prompt = _build_overview_prompt(overview_data, alert_samples, blocked_ips)

    try:
        result = call_agent(prompt, session_id=session_id)
    except AgentCallError as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    parsed = result.get("parsed") if isinstance(result.get("parsed"), dict) else {}
    return {
        "display": _build_overview_display(parsed, str(result.get("raw_text", ""))),
        "raw": parsed or {"raw_text": result.get("raw_text", "")},
        "usage": result.get("usage", {}),
        "session_id": result.get("session_id", ""),
        "request_id": result.get("request_id", ""),
    }


@app.get("/api/logs")
async def logs(
    request: Request,
    alerts_only: bool = False,
    traffic_kind: str | None = None,
    action: str | None = None,
    keyword: str | None = None,
    severity: str | None = None,
    alert_status: str | None = None,
    handled_status: str | None = None,
    page: int = 1,
    page_size: int = 20,
):
    require_api_auth(request)
    page = max(1, page)
    page_size = max(1, min(page_size, 100))
    return list_logs(
        page=page,
        page_size=page_size,
        alerts_only=alerts_only,
        traffic_kind=traffic_kind or None,
        action=action or None,
        keyword=keyword or None,
        severity=severity or None,
        alert_status=alert_status or None,
        handled_status=handled_status or None,
    )


@app.get("/api/logs/{log_id}")
async def log_detail(log_id: int, request: Request):
    require_api_auth(request)
    detail = get_log_detail(log_id)
    if not detail:
        raise HTTPException(status_code=404, detail="日志不存在")

    ip = detail.get("client_ip", "")
    geo = get_cached_ip_geo(ip)
    if not geo:
        geo = lookup_ip_geo(ip)
        if should_cache_geo_result(geo):
            cache_ip_geo(ip, geo)

    detail["ip_geo"] = geo
    return detail


@app.post("/api/agent/log/{log_id}/analyze")
async def agent_log_analyze(log_id: int, request: Request):
    require_api_auth(request)
    payload = await request.json() if request.headers.get("content-type", "").startswith("application/json") else {}
    session_id = str(payload.get("session_id", "")).strip() or None

    detail = get_log_detail(log_id)
    if not detail:
        raise HTTPException(status_code=404, detail="日志不存在")

    ip = detail.get("client_ip", "")
    geo = get_cached_ip_geo(ip)
    if not geo:
        geo = lookup_ip_geo(ip)
        if should_cache_geo_result(geo):
            cache_ip_geo(ip, geo)
    detail["ip_geo"] = geo

    prompt = _build_single_log_prompt(detail)
    try:
        result = call_agent(prompt, session_id=session_id)
    except AgentCallError as exc:
        raise HTTPException(status_code=502, detail=str(exc))

    parsed = result.get("parsed") if isinstance(result.get("parsed"), dict) else {}
    return {
        "display": _build_log_display(parsed, str(result.get("raw_text", ""))),
        "raw": parsed or {"raw_text": result.get("raw_text", "")},
        "usage": result.get("usage", {}),
        "session_id": result.get("session_id", ""),
        "request_id": result.get("request_id", ""),
    }


@app.patch("/api/logs/{log_id}/status")
async def patch_log_status(log_id: int, request: Request):
    require_api_auth(request)
    payload = await request.json()
    alert_status = str(payload.get("alert_status", "")).strip()

    if alert_status not in {"real_attack", "customer_business", "pending_business", "notified_event", "whitelist_traffic"}:
        raise HTTPException(status_code=400, detail="处置分类不合法")

    update_log_status(log_id, alert_status)
    return {"message": "ok"}


@app.post("/api/logs/disposition/bulk")
async def bulk_patch_log_status(request: Request):
    require_api_auth(request)
    payload = await request.json()
    alert_status = str(payload.get("alert_status", "")).strip()
    log_ids = payload.get("log_ids", [])

    if alert_status not in {"real_attack", "customer_business", "pending_business", "notified_event", "whitelist_traffic"}:
        raise HTTPException(status_code=400, detail="处置分类不合法")
    if not isinstance(log_ids, list) or not log_ids:
        raise HTTPException(status_code=400, detail="请选择需要处置的流量记录")

    bulk_update_log_status(log_ids, alert_status)
    return {"message": "ok"}


@app.get("/api/blocked-ips")
async def blocked_ips(
    request: Request,
    manual_page: int = 1,
    auto_page: int = 1,
    page_size: int = 20,
):
    require_api_auth(request)
    page_size = max(1, min(page_size, 100))
    manual_page = max(1, manual_page)
    auto_page = max(1, auto_page)
    manual = list_blocked_ips(page=manual_page, page_size=page_size)
    auto = list_cc_bans(page=auto_page, page_size=page_size)
    return {
        "manual": manual,
        "auto": auto,
        "counts": {
            "manual": manual["total"],
            "auto": auto["total"],
            "total": manual["total"] + auto["total"],
        },
    }


@app.post("/api/blocked-ips")
async def create_blocked_ip(request: Request):
    require_api_auth(request)
    payload = await request.json()
    ip = str(payload.get("ip", "")).strip()
    reason = str(payload.get("reason", "")).strip()

    if not ip:
        raise HTTPException(status_code=400, detail="IP 地址不能为空")

    add_blocked_ip(ip, reason or "手动封禁", created_by=request.session.get("username", "admin"))
    return {"message": "ok"}


@app.delete("/api/blocked-ips/{record_id}")
async def delete_blocked_ip(record_id: int, request: Request):
    require_api_auth(request)
    remove_blocked_ip(record_id)
    return {"message": "ok"}


@app.delete("/api/cc-bans/{record_id}")
async def delete_cc_ban(record_id: int, request: Request):
    require_api_auth(request)
    remove_cc_ban(record_id)
    return {"message": "ok"}


@app.get("/api/log-analysis/data")
async def log_analysis_data(request: Request, hours: int = 24):
    """获取用于日志分析的原始IP异常数据"""
    require_api_auth(request)
    hours = max(1, min(hours, 168))
    return get_ip_analysis_data(hours=hours)


@app.post("/api/agent/log-analysis")
async def agent_log_analysis(request: Request):
    """触发AI对24小时日志进行异常IP分析"""
    require_api_auth(request)
    payload = await request.json() if request.headers.get("content-type", "").startswith("application/json") else {}
    session_id = str(payload.get("session_id", "")).strip() or None
    hours = int(payload.get("hours", 24))
    hours = max(1, min(hours, 168))

    try:
        analysis_data = get_ip_analysis_data(hours=hours)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"获取分析数据失败: {str(exc)[:200]}")

    # 检查 AI 服务配置
    has_ai_config = bool(settings.dashscope_api_key and settings.bailian_app_id)

    if has_ai_config:
        prompt = _build_log_analysis_prompt(analysis_data)

        try:
            result = call_agent(prompt, session_id=session_id, timeout_seconds=max(300, settings.bailian_timeout))
        except AgentCallError as exc:
            # AI 服务失败，使用本地降级模式
            reason = f"AI 服务调用失败: {str(exc)[:100]}"
            display = _build_local_log_analysis_display(analysis_data, reason)
            return {
                "display": display,
                "raw": {"raw_text": reason},
                "analysis_time": analysis_data.get("analysis_time", ""),
                "mode": "local_fallback"
            }
        except Exception as exc:
            # 其他错误，使用本地降级模式
            reason = f"未知错误: {str(exc)[:100]}"
            display = _build_local_log_analysis_display(analysis_data, reason)
            return {
                "display": display,
                "raw": {"raw_text": reason},
                "analysis_time": analysis_data.get("analysis_time", ""),
                "mode": "local_fallback"
            }

        parsed = result.get("parsed") if isinstance(result.get("parsed"), dict) else {}
        raw_text = str(result.get("raw_text", ""))

        if not parsed and raw_text.strip():
            repair_prompt = _build_log_analysis_repair_prompt(raw_text)
            try:
                repaired = call_agent(repair_prompt, session_id=result.get("session_id") or session_id, timeout_seconds=120)
                repaired_parsed = repaired.get("parsed") if isinstance(repaired.get("parsed"), dict) else {}
                if repaired_parsed:
                    parsed = repaired_parsed
                    raw_text = str(repaired.get("raw_text", ""))
                    if not result.get("request_id") and repaired.get("request_id"):
                        result["request_id"] = repaired.get("request_id")
                    if not result.get("session_id") and repaired.get("session_id"):
                        result["session_id"] = repaired.get("session_id")
            except AgentCallError:
                # 修复失败，使用本地降级模式
                reason = "AI 结果解析失败，使用本地规则分析"
                display = _build_local_log_analysis_display(analysis_data, reason)
                return {
                    "display": display,
                    "raw": {"raw_text": reason},
                    "analysis_time": analysis_data.get("analysis_time", ""),
                    "mode": "local_fallback"
                }

        display = _build_log_analysis_display(parsed, raw_text, analysis_data)
        return {
            "display": display,
            "raw": parsed or {"raw_text": raw_text},
            "usage": result.get("usage", {}),
            "session_id": result.get("session_id", ""),
            "request_id": result.get("request_id", ""),
            "analysis_time": analysis_data.get("analysis_time", ""),
            "mode": "ai"
        }
    else:
        # 未配置 AI 服务，使用本地降级模式
        reason = "未配置 AI 服务，使用本地规则分析"
        display = _build_local_log_analysis_display(analysis_data, reason)
        return {
            "display": display,
            "raw": {"raw_text": reason},
            "analysis_time": analysis_data.get("analysis_time", ""),
            "mode": "local_fallback"
        }


def _set_doc_run_font(run) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def _add_doc_paragraph(doc, text: str, *, style: str | None = None):
    p = doc.add_paragraph(style=style) if style else doc.add_paragraph()
    run = p.add_run(str(text))
    _set_doc_run_font(run)
    return p


@app.post("/api/log-analysis/export")
async def export_log_analysis(request: Request):
    """将日志分析结果导出为 docx 文件"""
    require_api_auth(request)
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
    except ImportError:
        raise HTTPException(status_code=500, detail="python-docx 未安装，请执行 pip install python-docx")

    payload = await request.json() if request.headers.get("content-type", "").startswith("application/json") else {}
    display = _ensure_log_analysis_display_sections(dict(payload.get("display", {}) or {}))
    analysis_time = payload.get("analysis_time", datetime.now(timezone.utc).isoformat())

    globals()["qn"] = qn
    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Microsoft YaHei"
    normal_style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal_style.font.size = Pt(11)

    title_para = doc.add_heading(level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(display.get("title") or "24小时日志异常IP分析报告")
    title_run.bold = True
    title_run.font.size = Pt(24)
    _set_doc_run_font(title_run)

    try:
        dt = datetime.fromisoformat(analysis_time)
        time_str = dt.strftime("%Y年%m月%d日 %H:%M:%S UTC")
    except Exception:
        time_str = analysis_time
    meta_para = _add_doc_paragraph(doc, f"报告生成时间：{time_str}")
    meta_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    threat_map = {"critical": "严重", "high": "高危", "medium": "中危", "low": "低危"}
    threat_level = display.get("threat_level", "medium")
    tl_para = doc.add_paragraph()
    tl_run = tl_para.add_run(f"整体威胁等级：{threat_map.get(threat_level, threat_level)}")
    tl_run.bold = True
    tl_run.font.size = Pt(13)
    _set_doc_run_font(tl_run)
    if threat_level in ("critical", "high"):
        tl_run.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)
    elif threat_level == "medium":
        tl_run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    else:
        tl_run.font.color.rgb = RGBColor(0x16, 0x7A, 0x6E)

    overview = _normalize_paragraph_text(display.get("overview_summary", ""))
    if overview:
        doc.add_heading("一、整体概述", level=1)
        for line in [seg.strip() for seg in str(overview).splitlines() if seg.strip()]:
            _add_doc_paragraph(doc, line)

    key_findings = _normalize_list(display.get("key_findings", []))
    if key_findings:
        doc.add_heading("二、关键发现", level=1)
        for item in key_findings:
            _add_doc_paragraph(doc, str(item), style="List Bullet")

    attack_patterns = _normalize_list(display.get("attack_patterns", []))
    if attack_patterns:
        doc.add_heading("三、攻击模式分析", level=1)
        for item in attack_patterns:
            _add_doc_paragraph(doc, str(item), style="List Bullet")

    risk_summary = _normalize_paragraph_text(display.get("risk_summary", ""))
    if risk_summary:
        doc.add_heading("四、风险评估", level=1)
        _add_doc_paragraph(doc, risk_summary)

    anomaly_ips = _normalize_anomaly_ips(display.get("anomaly_ips", []))
    if anomaly_ips:
        doc.add_heading("五、异常IP详情", level=1)
        risk_label = {"critical": "严重", "high": "高危", "medium": "中危", "low": "低危"}
        for idx, ip_item in enumerate(anomaly_ips, 1):
            if not isinstance(ip_item, dict):
                continue
            sub = doc.add_heading(f"{idx}. {ip_item.get('ip', '未知IP')}", level=2)
            for run in sub.runs:
                _set_doc_run_font(run)
            rl = ip_item.get("risk_level", "medium")
            _add_doc_paragraph(doc, f"风险等级：{risk_label.get(rl, rl)}")
            anomaly_type = _normalize_paragraph_text(ip_item.get("anomaly_type"))
            evidence = _normalize_paragraph_text(ip_item.get("evidence"))
            if anomaly_type:
                _add_doc_paragraph(doc, f"异常类型：{anomaly_type}")
            if evidence:
                _add_doc_paragraph(doc, f"异常证据：{evidence}")
            ban = ip_item.get("ban_suggestion", {})
            if isinstance(ban, dict) and ban:
                days = ban.get("duration_days", "")
                reason = ban.get("reason", "")
                suggestion = f"建议封禁 {days} 天" if days else "建议封禁"
                if reason:
                    suggestion += f"，原因：{reason}"
                _add_doc_paragraph(doc, f"处置建议：{suggestion}")

    ban_policy = _normalize_paragraph_text(display.get("ban_policy_advice", ""))
    if ban_policy:
        doc.add_heading("六、IP封禁策略建议", level=1)
        _add_doc_paragraph(doc, ban_policy)

    recommendations = _normalize_list(display.get("recommendations", []))
    if recommendations:
        doc.add_heading("七、安全建议", level=1)
        for item in recommendations:
            _add_doc_paragraph(doc, str(item), style="List Bullet")

    security_notices = _normalize_list(display.get("security_notices", []))
    if security_notices:
        doc.add_heading("八、安全注意事项", level=1)
        for item in security_notices:
            _add_doc_paragraph(doc, str(item), style="List Bullet")

    next_steps = _normalize_list(display.get("next_steps", []))
    if next_steps:
        doc.add_heading("九、下一步行动", level=1)
        for item in next_steps:
            _add_doc_paragraph(doc, str(item), style="List Number")

    raw_data = display.get("raw_analysis_data", {})
    summary_data = raw_data.get("summary", {}) if isinstance(raw_data, dict) else {}
    if summary_data:
        appendix = doc.add_heading("附录：原始统计数据", level=1)
        for run in appendix.runs:
            _set_doc_run_font(run)
        table = doc.add_table(rows=1, cols=2)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "指标"
        hdr[1].text = "数值"
        label_map = {
            "total_requests": "24h 总请求数",
            "unique_ips": "独立源 IP 数",
            "blocked_count": "拦截次数",
            "high_count": "高危事件数",
            "brute_force_count": "暴力破解次数",
            "webshell_count": "WebShell 上传次数",
            "sql_injection_count": "SQL 注入次数",
            "cve_count": "CVE 利用次数",
            "scanner_count": "扫描探测次数",
        }
        for key, label in label_map.items():
            row = table.add_row().cells
            row[0].text = label
            row[1].text = str(summary_data.get(key, 0))
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        _set_doc_run_font(run)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            _set_doc_run_font(run)

    # 输出为字节流
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    try:
        dt = datetime.fromisoformat(analysis_time)
        filename_date = dt.strftime("%Y%m%d_%H%M")
    except Exception:
        filename_date = "report"

    filename = f"log_analysis_{filename_date}.docx"
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


if __name__ == "__main__":
    uvicorn.run("app.admin:app", host="0.0.0.0", port=9443)
